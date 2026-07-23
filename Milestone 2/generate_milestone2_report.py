#!/usr/bin/env python3
"""
Spider Text-to-SQL Deep SQL EDA report generator.

Uses only local artifacts in this directory:
  - schema.json
  - database/<db_id>/<db_id>.sqlite
  - *_gold.sql (e.g. dev_gold.sql)

All statistics are computed from data (nothing hardcoded into the report).
"""

from __future__ import annotations

import csv
import json
import os
import re
import sqlite3
import statistics
from collections import Counter, defaultdict
from pathlib import Path

ROOT = Path(__file__).resolve().parent
FIGURES = ROOT / "figures"
CSV_DIR = ROOT / "eda_csvs"
FIGURES.mkdir(exist_ok=True)
CSV_DIR.mkdir(exist_ok=True)

# Matplotlib wants a writable config dir in sandboxed/CI environments
os.environ.setdefault("MPLCONFIGDIR", str(ROOT / ".mplconfig"))
(ROOT / ".mplconfig").mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_json(path: Path):
    with open(path) as f:
        return json.load(f)


def classify_sql(query: str) -> dict:
    q = query.upper()
    return {
        "has_join": bool(re.search(r"\bJOIN\b", q)),
        "has_subquery": q.count("SELECT") > 1,
        "has_group_by": bool(re.search(r"\bGROUP\s+BY\b", q)),
        "has_order_by": bool(re.search(r"\bORDER\s+BY\b", q)),
        "has_having": bool(re.search(r"\bHAVING\b", q)),
        "has_limit": bool(re.search(r"\bLIMIT\b", q)),
        "has_distinct": bool(re.search(r"\bDISTINCT\b", q)),
        "has_union": bool(re.search(r"\bUNION\b", q)),
        "has_intersect": bool(re.search(r"\bINTERSECT\b", q)),
        "has_except": bool(re.search(r"\bEXCEPT\b", q)),
        "has_aggregation": bool(re.search(r"\b(COUNT|SUM|AVG|MIN|MAX)\s*\(", q)),
        "has_where": bool(re.search(r"\bWHERE\b", q)),
        "has_like": bool(re.search(r"\bLIKE\b", q)),
        "has_between": bool(re.search(r"\bBETWEEN\b", q)),
        "has_in": bool(re.search(r"\bIN\s*\(", q)),
    }


def count_joins(query: str) -> int:
    return len(re.findall(r"\bJOIN\b", query, re.I))


def agg_functions(query: str) -> list[str]:
    return [a.upper() for a in re.findall(r"\b(COUNT|SUM|AVG|MIN|MAX)\s*\(", query, re.I)]


def complexity_tier(query: str) -> str:
    f = classify_sql(query)
    score = sum([
        f["has_join"], f["has_subquery"], f["has_group_by"], f["has_having"],
        f["has_union"], f["has_intersect"], f["has_except"], count_joins(query) > 1,
    ])
    if score == 0:
        return "simple"
    if score == 1:
        return "moderate"
    if score <= 3:
        return "complex"
    return "extra"


def spider_hardness(query: str) -> str:
    """Approx Spider hardness (component-count style heuristic)."""
    q = query.upper()
    n_select = len(re.findall(r"\bSELECT\b", q))
    n_where = len(re.findall(r"\bWHERE\b", q))
    n_group = len(re.findall(r"\bGROUP\s+BY\b", q))
    n_order = len(re.findall(r"\bORDER\s+BY\b", q))
    n_join = count_joins(query)
    nested = n_select > 1
    setop = bool(re.search(r"\b(UNION|INTERSECT|EXCEPT)\b", q))
    components = n_where + n_group + n_order + n_join
    if nested or setop or components >= 4 or n_join >= 2:
        if nested and (setop or n_join >= 2 or components >= 3):
            return "extra hard"
        return "hard"
    if components >= 2 or n_join == 1 or re.search(r"\b(COUNT|SUM|AVG|MIN|MAX)\s*\(", q):
        return "medium"
    return "easy"


def categorize_domain(db_id: str) -> str:
    name = db_id.lower()
    rules = [
        ("education", ["student", "college", "school", "course", "university", "dorm"]),
        ("healthcare", ["hospital", "patient", "doctor", "disease", "clinic"]),
        ("sports", ["soccer", "basketball", "wta", "game", "player", "stadium", "baseball"]),
        ("entertainment", ["music", "movie", "singer", "concert", "tv", "orchestra", "cinema"]),
        ("business", ["store", "company", "employee", "hr", "customer", "invoice", "shop"]),
        ("transport", ["flight", "car", "bike", "ship", "aircraft", "train", "railway"]),
        ("geography", ["world", "geo", "country", "city", "address"]),
        ("food", ["restaurant", "wine", "food", "bakery"]),
        ("science", ["scientist", "protein", "planet", "gas"]),
    ]
    for domain, kws in rules:
        if any(k in name for k in kws):
            return domain
    return "other"


def write_csv(path: Path, rows: list[dict], fieldnames: list[str] | None = None):
    if not rows:
        path.write_text("")
        return
    fields = fieldnames or list(rows[0].keys())
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)


def _stats(vals: list) -> dict:
    if not vals:
        return {"min": 0, "max": 0, "mean": 0, "median": 0}
    return {
        "min": min(vals),
        "max": max(vals),
        "mean": round(statistics.mean(vals), 2),
        "median": statistics.median(vals),
    }


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_gold_sql_files() -> list[dict]:
    """Load all local *_gold.sql files as {query, db_id, split}."""
    examples = []
    for path in sorted(ROOT.glob("*_gold.sql")):
        split = path.stem.replace("_gold", "")  # train / dev / test
        with open(path, encoding="utf-8", errors="replace") as f:
            for line in f:
                line = line.strip()
                if not line or "\t" not in line:
                    continue
                query, db_id = line.rsplit("\t", 1)
                examples.append({"query": query.strip(), "db_id": db_id.strip(), "split": split})
    return examples


def load_schemas() -> list[dict]:
    path = ROOT / "schema.json"
    if not path.exists():
        raise FileNotFoundError(f"Missing {path}")
    return load_json(path)


# ---------------------------------------------------------------------------
# Schema helpers (Spider 1-based column indices, index 0 was '*')
# ---------------------------------------------------------------------------

def build_schema_index(schemas: list[dict]) -> dict[str, dict]:
    """
    Rebuild column→table mapping by walking SQLite in schema table_names order.
    Falls back to equal-split heuristic if sqlite is missing.
    """
    index = {}
    for s in schemas:
        db_id = s["db_id"]
        db_path = ROOT / "database" / db_id / f"{db_id}.sqlite"
        col_to_table: dict[str, str] = {}
        ordered_cols: list[tuple[str, str]] = []  # (table, col) in spider export order
        table_cols: dict[str, list[str]] = {}

        if db_path.exists():
            conn = sqlite3.connect(db_path)
            try:
                for tname in s["table_names"]:
                    try:
                        info = conn.execute(f'PRAGMA table_info("{tname}")').fetchall()
                    except sqlite3.Error:
                        # try case-insensitive match
                        real = None
                        for (name,) in conn.execute(
                            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
                        ):
                            if name.lower() == tname.lower():
                                real = name
                                break
                        if real is None:
                            info = []
                        else:
                            info = conn.execute(f'PRAGMA table_info("{real}")').fetchall()
                            tname = real
                    cols = [c[1] for c in info]
                    table_cols[tname] = cols
                    for c in cols:
                        ordered_cols.append((tname, c))
                        col_to_table[f"{tname}.{c}".lower()] = tname
                        col_to_table[c.lower()] = tname  # last-writer wins for bare names
            finally:
                conn.close()
        else:
            # Fallback: use schema column list without table membership
            for c in s["column_names"]:
                ordered_cols.append(("", c))

        # Spider PK/FK indices are 1-based into original column list (0='*')
        def col_at(spider_idx: int) -> tuple[str, str] | None:
            i = spider_idx - 1
            if 0 <= i < len(s["column_names"]):
                name = s["column_names"][i]
                # Prefer ordered sqlite mapping if lengths match
                if len(ordered_cols) == len(s["column_names"]):
                    return ordered_cols[i]
                # else search by name
                for t, c in ordered_cols:
                    if c.lower() == name.lower():
                        return (t, c)
                return ("", name)
            return None

        fks = []
        for a, b in s.get("foreign_keys", []):
            ca, cb = col_at(a), col_at(b)
            if ca and cb:
                fks.append({
                    "from_table": ca[0], "from_col": ca[1],
                    "to_table": cb[0], "to_col": cb[1],
                    "from_idx": a, "to_idx": b,
                })

        pks = []
        for pk in s.get("primary_keys", []):
            c = col_at(pk)
            if c:
                pks.append({"table": c[0], "column": c[1], "idx": pk})

        index[db_id] = {
            "schema": s,
            "table_cols": table_cols,
            "fks": fks,
            "pks": pks,
            "ordered_cols": ordered_cols,
            "fk_pairs": {
                frozenset({
                    (fk["from_table"].lower(), fk["from_col"].lower()),
                    (fk["to_table"].lower(), fk["to_col"].lower()),
                })
                for fk in fks
                if fk["from_table"] and fk["to_table"]
            },
            "fk_col_pairs": {
                frozenset({fk["from_col"].lower(), fk["to_col"].lower()})
                for fk in fks
            },
        }
    return index


# ---------------------------------------------------------------------------
# Database / table / column profiling
# ---------------------------------------------------------------------------

def profile_databases(schemas: list[dict]) -> tuple[list[dict], list[dict], list[dict]]:
    """Return (db_profiles, table_rows, column_rows)."""
    db_profiles, table_rows, column_rows = [], [], []

    for s in schemas:
        db_id = s["db_id"]
        db_path = ROOT / "database" / db_id / f"{db_id}.sqlite"
        domain = categorize_domain(db_id)
        if not db_path.exists():
            db_profiles.append({
                "db_id": db_id, "tables": s["n_tables"], "rows": 0, "columns": s["n_columns"],
                "null_pct": 0.0, "domain": domain, "sqlite_found": 0,
                "fk_count": len(s.get("foreign_keys", [])), "pk_count": len(s.get("primary_keys", [])),
            })
            continue

        conn = sqlite3.connect(db_path)
        try:
            tbls = [
                r[0] for r in conn.execute(
                    "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
                )
            ]
            total_rows = total_cols = total_nulls = total_cells = 0
            for t in tbls:
                info = conn.execute(f'PRAGMA table_info("{t}")').fetchall()
                try:
                    nrow = conn.execute(f'SELECT COUNT(*) FROM "{t}"').fetchone()[0]
                except sqlite3.Error:
                    nrow = 0
                fks = conn.execute(f'PRAGMA foreign_key_list("{t}")').fetchall()
                pk_cols = {c[1] for c in info if c[5] > 0}
                total_rows += nrow
                total_cols += len(info)

                table_nulls = 0
                for c in info:
                    col = c[1]
                    ctype = c[2] or ""
                    try:
                        nulls = conn.execute(
                            f'SELECT COUNT(*) FROM "{t}" WHERE "{col}" IS NULL'
                        ).fetchone()[0]
                        ndistinct = conn.execute(
                            f'SELECT COUNT(DISTINCT "{col}") FROM "{t}"'
                        ).fetchone()[0]
                    except sqlite3.Error:
                        nulls, ndistinct = 0, 0
                    table_nulls += nulls
                    total_nulls += nulls
                    total_cells += nrow
                    column_rows.append({
                        "db_id": db_id,
                        "table": t,
                        "column": col,
                        "sqlite_type": ctype,
                        "row_count": nrow,
                        "distinct_count": ndistinct,
                        "null_count": nulls,
                        "null_pct": round(100 * nulls / nrow, 2) if nrow else 0.0,
                        "is_pk": int(col in pk_cols),
                        "cardinality_ratio": round(ndistinct / nrow, 4) if nrow else 0.0,
                    })

                table_rows.append({
                    "db_id": db_id,
                    "table": t,
                    "n_columns": len(info),
                    "n_rows": nrow,
                    "n_fk_declared": len(fks),
                    "n_pk_cols": len(pk_cols),
                    "null_cells": table_nulls,
                    "null_pct": round(100 * table_nulls / (nrow * len(info)), 2) if nrow and info else 0.0,
                    "domain": domain,
                })

            db_profiles.append({
                "db_id": db_id,
                "tables": len(tbls),
                "rows": total_rows,
                "columns": total_cols,
                "null_pct": round(100 * total_nulls / total_cells, 2) if total_cells else 0.0,
                "domain": domain,
                "sqlite_found": 1,
                "fk_count": len(s.get("foreign_keys", [])),
                "pk_count": len(s.get("primary_keys", [])),
            })
        finally:
            conn.close()

    return db_profiles, table_rows, column_rows


def schema_level_stats(schemas: list[dict]) -> dict:
    t_cnt = [s["n_tables"] for s in schemas]
    c_cnt = [s["n_columns"] for s in schemas]
    fk_cnt = [len(s.get("foreign_keys", [])) for s in schemas]
    col_types = Counter()
    for s in schemas:
        for ct in s.get("column_types", []):
            col_types[ct] += 1
    return {
        "n_dbs": len(schemas),
        "tables": _stats(t_cnt),
        "columns": _stats(c_cnt),
        "foreign_keys": _stats(fk_cnt),
        "column_types": dict(col_types.most_common()),
    }


# ---------------------------------------------------------------------------
# SQL parsing for join / relation analysis
# ---------------------------------------------------------------------------

_JOIN_TYPE_RE = re.compile(
    r"\b((?:INNER\s+|LEFT\s+(?:OUTER\s+)?|RIGHT\s+(?:OUTER\s+)?|FULL\s+(?:OUTER\s+)?|CROSS\s+)?JOIN)\b",
    re.I,
)
_TABLE_REF_RE = re.compile(
    r"\b(?:FROM|JOIN)\s+([A-Za-z_][\w]*)(?:\s+(?:AS\s+)?([A-Za-z_][\w]*))?",
    re.I,
)
_ON_COND_RE = re.compile(
    r"\bON\s+([A-Za-z_][\w]*)\.([A-Za-z_][\w]*)\s*=\s*([A-Za-z_][\w]*)\.([A-Za-z_][\w]*)",
    re.I,
)
# Implicit comma joins: FROM a, b WHERE a.x = b.y — softer signal
_WHERE_EQ_RE = re.compile(
    r"\b([A-Za-z_][\w]*)\.([A-Za-z_][\w]*)\s*=\s*([A-Za-z_][\w]*)\.([A-Za-z_][\w]*)",
    re.I,
)
_WHERE_OPS = [
    ("=", r"(?<![<>!=])=(?!=)"),
    (">", r"(?<![<>])>(?!=)"),
    ("<", r"(?<![<>])<(?!=)"),
    (">=", r">="),
    ("<=", r"<="),
    ("!=/<>", r"(?:!=|<>)"),
    ("LIKE", r"\bLIKE\b"),
    ("IN", r"\bIN\s*\("),
    ("BETWEEN", r"\bBETWEEN\b"),
    ("IS NULL", r"\bIS\s+NULL\b"),
]


def extract_table_refs(query: str) -> list[str]:
    tables = []
    for m in _TABLE_REF_RE.finditer(query):
        tables.append(m.group(1))
    # unique preserve order
    seen, out = set(), []
    for t in tables:
        key = t.lower()
        if key not in seen:
            seen.add(key)
            out.append(t)
    return out


def extract_join_types(query: str) -> list[str]:
    types = []
    for m in _JOIN_TYPE_RE.finditer(query):
        raw = re.sub(r"\s+", " ", m.group(1).upper()).strip()
        if raw == "JOIN":
            types.append("INNER")  # bare JOIN == INNER in Spider
        elif raw.startswith("INNER"):
            types.append("INNER")
        elif raw.startswith("LEFT"):
            types.append("LEFT")
        elif raw.startswith("RIGHT"):
            types.append("RIGHT")
        elif raw.startswith("FULL"):
            types.append("FULL")
        elif raw.startswith("CROSS"):
            types.append("CROSS")
        else:
            types.append(raw)
    return types


def extract_join_conditions(query: str) -> list[tuple[str, str, str, str]]:
    """Return list of (alias1, col1, alias2, col2)."""
    return [(a, b, c, d) for a, b, c, d in _ON_COND_RE.findall(query)]


def subquery_locations(query: str) -> list[str]:
    """Heuristic: classify nested SELECT by preceding clause keyword."""
    locs = []
    for m in re.finditer(r"\bSELECT\b", query, re.I):
        if m.start() == re.search(r"\bSELECT\b", query, re.I).start():
            continue  # top-level
        prefix = query[: m.start()].upper()
        # last significant keyword
        last = None
        for kw in ("WHERE", "HAVING", "FROM", "SELECT", "AND", "OR", "ON"):
            idx = prefix.rfind(kw)
            if idx >= 0 and (last is None or idx > last[0]):
                last = (idx, kw)
        locs.append((last[1] if last else "OTHER"))
    return locs


def relation_topology(n_tables: int, n_joins: int) -> str:
    if n_tables <= 1:
        return "single"
    if n_tables == 2:
        return "binary"
    if n_joins >= n_tables:  # more edges than tree → mesh-ish
        return "mesh"
    if n_tables >= 4 and n_joins == n_tables - 1:
        # cannot distinguish star vs chain without graph; use join count heuristic
        return "chain_or_star"
    if n_joins == n_tables - 1:
        return "chain"
    if n_joins < n_tables - 1:
        return "partial"
    return "mesh"


def analyze_sql_examples(examples: list[dict], schema_index: dict) -> dict:
    by_split: dict[str, list] = defaultdict(list)
    for e in examples:
        by_split[e["split"]].append(e)

    per_query_rows = []
    join_pair_rows = []
    join_type_counter = Counter()
    where_op_counter = Counter()
    agg_counter = Counter()
    hardness_counter = Counter()
    topology_counter = Counter()
    subquery_loc_counter = Counter()
    table_ref_dist = Counter()
    join_count_dist = Counter()
    fk_aligned = fk_total = 0
    db_join_stats = defaultdict(lambda: {"n": 0, "with_join": 0})

    for e in examples:
        q = e["query"]
        db_id = e["db_id"]
        feats = classify_sql(q)
        n_joins = count_joins(q)
        tables = extract_table_refs(q)
        jtypes = extract_join_types(q)
        join_type_counter.update(jtypes)
        join_count_dist[n_joins] += 1
        table_ref_dist[len(tables)] += 1
        aggs = agg_functions(q)
        agg_counter.update(aggs)
        hard = spider_hardness(q)
        hardness_counter[hard] += 1
        topo = relation_topology(len(tables), n_joins)
        topology_counter[topo] += 1
        for loc in subquery_locations(q):
            subquery_loc_counter[loc] += 1

        # WHERE operators (only inside WHERE… before GROUP/ORDER/LIMIT)
        where_m = re.search(r"\bWHERE\b(.*?)(?:\bGROUP\s+BY\b|\bORDER\s+BY\b|\bLIMIT\b|\bHAVING\b|$)", q, re.I | re.S)
        if where_m:
            clause = where_m.group(1)
            for name, pat in _WHERE_OPS:
                if re.search(pat, clause, re.I):
                    where_op_counter[name] += 1

        db_join_stats[db_id]["n"] += 1
        if feats["has_join"]:
            db_join_stats[db_id]["with_join"] += 1

        # FK alignment on ON conditions
        conds = extract_join_conditions(q)
        # Build alias map
        alias_map = {}
        for m in _TABLE_REF_RE.finditer(q):
            t, alias = m.group(1), m.group(2)
            alias_map[(alias or t).lower()] = t.lower()

        sx = schema_index.get(db_id, {})
        fk_pairs = sx.get("fk_pairs", set())
        fk_col_pairs = sx.get("fk_col_pairs", set())
        n_fk_match = 0
        for a1, c1, a2, c2 in conds:
            t1 = alias_map.get(a1.lower(), a1.lower())
            t2 = alias_map.get(a2.lower(), a2.lower())
            pair = frozenset({(t1, c1.lower()), (t2, c2.lower())})
            colpair = frozenset({c1.lower(), c2.lower()})
            matched = pair in fk_pairs or colpair in fk_col_pairs
            if matched:
                n_fk_match += 1
            fk_total += 1
            if matched:
                fk_aligned += 1
            join_pair_rows.append({
                "split": e["split"],
                "db_id": db_id,
                "left_table": t1,
                "left_col": c1,
                "right_table": t2,
                "right_col": c2,
                "fk_aligned": int(matched),
                "query": q[:200],
            })

        # Table co-occurrence pairs
        for i in range(len(tables)):
            for j in range(i + 1, len(tables)):
                a, b = sorted([tables[i].lower(), tables[j].lower()])
                join_pair_rows  # stored separately below

        per_query_rows.append({
            "split": e["split"],
            "db_id": db_id,
            "n_joins": n_joins,
            "n_tables_ref": len(tables),
            "tables": "|".join(tables),
            "join_types": "|".join(jtypes) if jtypes else "",
            "topology": topo,
            "hardness": hard,
            "tier": complexity_tier(q),
            "n_fk_join_conds": n_fk_match,
            "n_join_conds": len(conds),
            "has_subquery": int(feats["has_subquery"]),
            "has_group_by": int(feats["has_group_by"]),
            "has_aggregation": int(feats["has_aggregation"]),
            "has_where": int(feats["has_where"]),
            "agg_funcs": "|".join(aggs),
            "sql_tokens": len(q.split()),
            "query": q,
        })

    # Table co-occurrence from per-query tables
    pair_counter = Counter()
    for row in per_query_rows:
        ts = [t for t in row["tables"].split("|") if t]
        for i in range(len(ts)):
            for j in range(i + 1, len(ts)):
                a, b = sorted([ts[i].lower(), ts[j].lower()])
                pair_counter[(row["db_id"], a, b)] += 1

    def analyze_split(name: str, items: list[dict]) -> dict:
        if not items:
            return {"split": name, "n_examples": 0}
        queries = [x["query"] for x in items]
        feats = [classify_sql(q) for q in queries]
        tiers = Counter(complexity_tier(q) for q in queries)
        sql_tokens = [len(q.split()) for q in queries]
        feat_pct = {
            k: round(100 * sum(f[k] for f in feats) / len(feats), 1)
            for k in feats[0]
        }
        return {
            "split": name,
            "n_examples": len(items),
            "n_unique_dbs": len({x["db_id"] for x in items}),
            "sql_tokens": _stats(sql_tokens),
            "sql_complexity_pct": feat_pct,
            "complexity_tiers": {
                k: {"count": tiers[k], "pct": round(100 * tiers[k] / len(items), 1)}
                for k in ("simple", "moderate", "complex", "extra")
            },
            "join_distribution": dict(sorted(Counter(count_joins(q) for q in queries).items())),
            "hardness": dict(Counter(spider_hardness(q) for q in queries).most_common()),
            "aggregation_functions": dict(Counter(
                a for q in queries for a in agg_functions(q)
            ).most_common()),
        }

    splits = {name: analyze_split(name, items) for name, items in by_split.items()}
    # Combined
    splits["all"] = analyze_split("all", examples)

    db_join_rate = []
    for db_id, st in sorted(db_join_stats.items(), key=lambda x: -x[1]["with_join"] / max(x[1]["n"], 1)):
        db_join_rate.append({
            "db_id": db_id,
            "n_queries": st["n"],
            "n_with_join": st["with_join"],
            "join_rate_pct": round(100 * st["with_join"] / st["n"], 1) if st["n"] else 0,
        })

    table_pair_rows = [
        {"db_id": db, "table_a": a, "table_b": b, "cooccurrence": n}
        for (db, a, b), n in pair_counter.most_common()
    ]

    return {
        "splits": splits,
        "n_examples_total": len(examples),
        "join_types": dict(join_type_counter.most_common()),
        "where_operators": dict(where_op_counter.most_common()),
        "aggregations": dict(agg_counter.most_common()),
        "hardness": dict(hardness_counter.most_common()),
        "topology": dict(topology_counter.most_common()),
        "subquery_locations": dict(subquery_loc_counter.most_common()),
        "table_ref_distribution": dict(sorted(table_ref_dist.items())),
        "join_count_distribution": dict(sorted(join_count_dist.items())),
        "fk_alignment": {
            "join_conditions": fk_total,
            "fk_aligned": fk_aligned,
            "pct": round(100 * fk_aligned / fk_total, 1) if fk_total else 0.0,
        },
        "db_join_rates": db_join_rate,
        "per_query_rows": per_query_rows,
        "join_condition_rows": join_pair_rows,
        "table_pair_rows": table_pair_rows,
    }


def analyze_schema_joins(schema_index: dict) -> list[dict]:
    """Per-DB FK graph stats (schema-level join potential)."""
    rows = []
    for db_id, sx in schema_index.items():
        s = sx["schema"]
        n_tables = s["n_tables"]
        n_fks = len(sx["fks"])
        tables_with_fk = {fk["from_table"] for fk in sx["fks"]} | {fk["to_table"] for fk in sx["fks"]}
        rows.append({
            "db_id": db_id,
            "n_tables": n_tables,
            "n_columns": s["n_columns"],
            "n_foreign_keys": n_fks,
            "n_primary_keys": len(sx["pks"]),
            "tables_in_fk_graph": len(tables_with_fk),
            "fk_coverage_pct": round(100 * len(tables_with_fk) / n_tables, 1) if n_tables else 0,
            "domain": categorize_domain(db_id),
        })
    return rows


# ---------------------------------------------------------------------------
# Charts
# ---------------------------------------------------------------------------

def make_charts(eda: dict, sql_eda: dict) -> list[str]:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    saved = []

    def save(fig, name):
        fig.tight_layout()
        fig.savefig(FIGURES / name, dpi=150)
        plt.close(fig)
        saved.append(name)

    # Split sizes
    splits = [k for k in ("train", "dev", "test", "all") if k in sql_eda["splits"] and sql_eda["splits"][k]["n_examples"]]
    if not splits:
        splits = [k for k, v in sql_eda["splits"].items() if v.get("n_examples")]
    labels = [s.title() for s in splits]
    sizes = [sql_eda["splits"][s]["n_examples"] for s in splits]
    fig, ax = plt.subplots(figsize=(7, 4))
    bars = ax.bar(labels, sizes, color="#2563eb", edgecolor="white")
    ax.set_title("Examples per Split (local gold SQL)")
    ax.set_ylabel("Count")
    for b, v in zip(bars, sizes):
        ax.text(b.get_x() + b.get_width() / 2, v, str(v), ha="center", va="bottom", fontsize=10)
    save(fig, "split_sizes.png")

    # SQL complexity (prefer 'all', else first split)
    key = "all" if "all" in sql_eda["splits"] else splits[0]
    tc = sql_eda["splits"][key]["sql_complexity_pct"]
    keys = ["has_where", "has_aggregation", "has_join", "has_group_by",
            "has_order_by", "has_subquery", "has_limit", "has_distinct"]
    labels2 = ["WHERE", "Agg", "JOIN", "GROUP BY", "ORDER BY", "Subquery", "LIMIT", "DISTINCT"]
    vals = [tc.get(k, 0) for k in keys]
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.barh(labels2, vals, color="#2563eb", edgecolor="white")
    ax.set_xlabel("% of Queries")
    ax.set_title(f"SQL Feature Prevalence ({key})")
    for i, v in enumerate(vals):
        ax.text(v + 0.5, i, f"{v}%", va="center", fontsize=9)
    save(fig, "sql_complexity.png")

    # Complexity tiers
    fig, ax = plt.subplots(figsize=(8, 4))
    tier_keys = ["simple", "moderate", "complex", "extra"]
    tier_labels = ["Simple", "Moderate", "Complex", "Extra Hard"]
    tier_colors = ["#86efac", "#fde047", "#fb923c", "#ef4444"]
    plot_splits = [s for s in splits if s != "all"] or splits
    x = range(len(plot_splits))
    width = 0.2
    for i, (tk, tl, c) in enumerate(zip(tier_keys, tier_labels, tier_colors)):
        vals = [sql_eda["splits"][s]["complexity_tiers"][tk]["pct"] for s in plot_splits]
        ax.bar([xi + i * width for xi in x], vals, width, label=tl, color=c)
    ax.set_xticks([xi + 1.5 * width for xi in x])
    ax.set_xticklabels([s.title() for s in plot_splits])
    ax.set_ylabel("% of Queries")
    ax.set_title("Query Complexity Tiers by Split")
    ax.legend(fontsize=8)
    save(fig, "complexity_tiers.png")

    # Column types
    ct = eda["schema_stats"]["column_types"]
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.pie(list(ct.values()), labels=list(ct.keys()), autopct="%1.1f%%", startangle=140)
    ax.set_title("Column Data Types (schema.json)")
    save(fig, "column_types.png")

    # DB size
    rows = [p["rows"] for p in eda["database_profiles"] if p["rows"] > 0]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist(rows, bins=30, color="#0891b2", edgecolor="white")
    ax.set_xlabel("Total Rows per Database")
    ax.set_ylabel("Number of Databases")
    ax.set_title("Database Size Distribution")
    ax.set_yscale("log")
    save(fig, "db_size_distribution.png")

    # JOIN distribution
    jd = sql_eda["join_count_distribution"]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar([str(k) for k in jd], list(jd.values()), color="#ea580c", edgecolor="white")
    ax.set_xlabel("Number of JOINs per Query")
    ax.set_ylabel("Query Count")
    ax.set_title("JOIN Count Distribution")
    save(fig, "join_distribution.png")

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar([str(k) for k in jd], list(jd.values()), color="#ea580c", edgecolor="white")
    ax.set_xlabel("JOINs per Query")
    ax.set_ylabel("Count")
    ax.set_title("Deep EDA: JOIN Count")
    save(fig, "deep_join_count.png")

    # Join types
    jt = sql_eda["join_types"]
    fig, ax = plt.subplots(figsize=(7, 4))
    if jt:
        ax.bar(list(jt.keys()), list(jt.values()), color="#7c3aed", edgecolor="white")
    ax.set_title("JOIN Types")
    ax.set_ylabel("Clause Count")
    save(fig, "deep_join_types.png")

    # Hardness
    hd = sql_eda["hardness"]
    fig, ax = plt.subplots(figsize=(7, 4))
    order = ["easy", "medium", "hard", "extra hard"]
    labs = [k for k in order if k in hd] + [k for k in hd if k not in order]
    ax.bar(labs, [hd[k] for k in labs], color="#0ea5e9", edgecolor="white")
    ax.set_title("Spider Hardness (heuristic)")
    ax.set_ylabel("Count")
    save(fig, "deep_spider_hardness.png")

    # Hardness by split
    fig, ax = plt.subplots(figsize=(8, 4))
    for hi, h in enumerate(order):
        vals = []
        for s in plot_splits:
            vals.append(sql_eda["splits"][s]["hardness"].get(h, 0))
        ax.bar([xi + hi * 0.2 for xi in range(len(plot_splits))], vals, 0.2, label=h)
    ax.set_xticks([xi + 0.3 for xi in range(len(plot_splits))])
    ax.set_xticklabels([s.title() for s in plot_splits])
    ax.legend(fontsize=8)
    ax.set_title("Hardness by Split")
    save(fig, "deep_hardness_splits.png")

    # Table refs
    tr = sql_eda["table_ref_distribution"]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar([str(k) for k in tr], list(tr.values()), color="#16a34a", edgecolor="white")
    ax.set_xlabel("Tables referenced")
    ax.set_ylabel("Queries")
    ax.set_title("Table References per Query")
    save(fig, "deep_table_refs.png")

    # WHERE ops
    wo = sql_eda["where_operators"]
    fig, ax = plt.subplots(figsize=(8, 4))
    if wo:
        ax.barh(list(wo.keys()), list(wo.values()), color="#f59e0b", edgecolor="white")
    ax.set_title("WHERE Operator Frequency")
    save(fig, "deep_where_operators.png")

    # Aggregations
    ag = sql_eda["aggregations"]
    fig, ax = plt.subplots(figsize=(7, 4))
    if ag:
        ax.bar(list(ag.keys()), list(ag.values()), color="#dc2626", edgecolor="white")
    ax.set_title("Aggregation Functions")
    save(fig, "deep_aggregations.png")

    # Clause prevalence
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.barh(labels2, vals, color="#2563eb", edgecolor="white")
    ax.set_title("Clause Prevalence")
    save(fig, "deep_clause_prevalence.png")

    # Topology
    tp = sql_eda["topology"]
    fig, ax = plt.subplots(figsize=(8, 4))
    if tp:
        ax.bar(list(tp.keys()), list(tp.values()), color="#0891b2", edgecolor="white")
    ax.set_title("Query Relation Topology")
    plt.xticks(rotation=20, ha="right")
    save(fig, "deep_relation_topology.png")

    # FK alignment
    fa = sql_eda["fk_alignment"]
    fig, ax = plt.subplots(figsize=(5, 4))
    ax.bar(["FK-aligned", "Other"], [fa["fk_aligned"], max(fa["join_conditions"] - fa["fk_aligned"], 0)],
           color=["#16a34a", "#94a3b8"], edgecolor="white")
    ax.set_title(f"JOIN↔FK Alignment ({fa['pct']}%)")
    save(fig, "deep_fk_alignment.png")

    # Subquery types
    sq = sql_eda["subquery_locations"]
    fig, ax = plt.subplots(figsize=(7, 4))
    if sq:
        ax.bar(list(sq.keys()), list(sq.values()), color="#a855f7", edgecolor="white")
    ax.set_title("Subquery Locations")
    save(fig, "deep_subquery_types.png")

    # Top table pairs
    pairs = sql_eda["table_pair_rows"][:15]
    fig, ax = plt.subplots(figsize=(9, 5))
    if pairs:
        labels_p = [f"{r['db_id'][:10]}:{r['table_a'][:8]}-{r['table_b'][:8]}" for r in pairs]
        ax.barh(labels_p[::-1], [r["cooccurrence"] for r in pairs][::-1], color="#ea580c")
    ax.set_title("Top Table Co-occurrence Pairs")
    save(fig, "deep_table_pairs.png")

    # Tables per DB / rows per table / columns per table
    profiles = eda["database_profiles"]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist([p["tables"] for p in profiles], bins=20, color="#2563eb", edgecolor="white")
    ax.set_title("Tables per Database")
    ax.set_xlabel("Tables")
    save(fig, "tables_per_db.png")

    trows = eda["table_analysis"]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist([r["n_rows"] for r in trows if r["n_rows"] > 0], bins=40, color="#0891b2", edgecolor="white")
    ax.set_yscale("log")
    ax.set_title("Rows per Table")
    save(fig, "rows_per_table.png")

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist([r["n_columns"] for r in trows], bins=20, color="#7c3aed", edgecolor="white")
    ax.set_title("Columns per Table")
    save(fig, "columns_per_table.png")

    # DB join rates chart
    top = sql_eda["db_join_rates"][:12]
    fig, ax = plt.subplots(figsize=(9, 4))
    if top:
        ax.barh([r["db_id"] for r in top][::-1], [r["join_rate_pct"] for r in top][::-1], color="#dc2626")
    ax.set_xlabel("JOIN rate %")
    ax.set_title("Databases with Highest JOIN Rates")
    save(fig, "deep_question_join_rate.png")

    # Placeholder-compatible charts that existed before
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(list(tp.keys()), list(tp.values()), color="#0ea5e9")
    ax.set_title("Question→Topology proxy (by query shape)")
    plt.xticks(rotation=20, ha="right")
    save(fig, "deep_question_topology.png")

    # question_patterns — gold SQL has no NL; use hardness as stand-in chart already made
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(labs, [hd[k] for k in labs], color="#7c3aed")
    ax.set_title("Query Pattern Proxy (hardness)")
    save(fig, "question_patterns.png")

    return saved


# ---------------------------------------------------------------------------
# Report writing
# ---------------------------------------------------------------------------

def write_docx(eda: dict, sql_eda: dict, chart_files: list[str]):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    schema = eda["schema_stats"]
    dq = eda["data_quality"]
    all_split = sql_eda["splits"].get("all") or next(iter(sql_eda["splits"].values()))
    tc = all_split.get("sql_complexity_pct", {})
    tiers = all_split.get("complexity_tiers", {})

    def heading(text, level=1):
        doc.add_heading(text, level=level)

    def para(text):
        doc.add_paragraph(text)

    def bullets(items):
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

    def table(headers, rows):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = str(h)
        for ri, row in enumerate(rows, 1):
            for ci, val in enumerate(row):
                t.rows[ri].cells[ci].text = str(val)
        doc.add_paragraph("")

    def fig(filename, caption, width=5.8):
        path = FIGURES / filename
        if path.exists():
            doc.add_picture(str(path), width=Inches(width))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)

    # Cover
    title = doc.add_heading("Deep SQL EDA Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Spider Text-to-SQL Dataset: Complete EDA & Profiling Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph(
        f"Generated from local artifacts only — {eda['inputs']['n_schemas']} schemas, "
        f"{eda['inputs']['n_sqlite']} SQLite DBs, {sql_eda['n_examples_total']} gold SQL examples "
        f"({', '.join(eda['inputs']['gold_files']) or 'none'})."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1
    heading("1. Introduction")
    para(
        "This report documents exploratory data analysis (EDA) and database profiling of the local "
        "Spider Text-to-SQL artifacts. Given a natural language question and a relational schema, "
        "the task is to generate executable SQL. Analysis focuses on SQL structural complexity, "
        "schema linking difficulty (JOINs / foreign keys), table/column distributions, and data quality."
    )
    para(
        "All tables, charts, and percentages below are computed by generate_milestone2_report.py from "
        "schema.json, database/*.sqlite, and local *_gold.sql files — not handwritten statistics."
    )

    # 2
    heading("2. Objectives")
    bullets([
        "Validate local Spider schemas and SQLite databases",
        "Profile every database/table/column (rows, NULL rates, cardinality)",
        "Quantify SQL complexity from available gold SQL",
        "Analyze JOIN types, FK alignment, and relation topology",
        "Export reproducible CSV finding tables + charts + this document",
    ])

    # 3
    heading("3. Dataset Identification")
    para("Dataset: Spider — Large-Scale Human-Labeled Dataset for Complex and Cross-Domain Text-to-SQL (Yu et al., EMNLP 2018).")
    para("Local scope: train+dev style SQLite folder + schema.json + gold SQL present in this directory.")
    para("Task scope: Text-to-SQL only.")

    heading("3.1 Local Inputs Used", 2)
    table(["Artifact", "Role", "Count"], [
        ["schema.json", "Schema metadata (tables, columns, types, PK/FK)", eda["inputs"]["n_schemas"]],
        ["database/", "Executable SQLite databases", eda["inputs"]["n_sqlite"]],
        ["*_gold.sql", "Gold SQL queries for complexity / JOIN EDA", sql_eda["n_examples_total"]],
        ["eda_csvs/", "Exported finding tables (generated)", "see §16"],
    ])

    # 4
    heading("4. Dataset Structure")
    table(["File / Folder", "Role"], [
        ["schema.json", "Normalized schema metadata for train+dev DBs"],
        ["database/<db_id>/<db_id>.sqlite", "Executable DBs for profiling & execution eval"],
        ["dev_gold.sql (and any other *_gold.sql)", "query \\t db_id gold labels"],
        ["eda_csvs/*.csv", "Generated table & JOIN analysis exports"],
        ["figures/*.png", "Generated charts"],
    ])

    # 5
    heading("5. Gold SQL Record Format")
    para("Each gold SQL line is tab-separated: <SQLQuery>\\t<db_id>")

    # 6
    heading("6. Split Analysis (from local gold SQL)")
    fig("split_sizes.png", "Figure 1: Examples per available gold split")
    split_rows = []
    for name, st in sql_eda["splits"].items():
        if name == "all" or not st.get("n_examples"):
            continue
        split_rows.append([
            name, st["n_examples"], st["n_unique_dbs"],
            st["sql_tokens"]["mean"], st["sql_tokens"]["median"],
        ])
    table(["Split", "Examples", "Unique DBs", "Avg SQL tokens", "Median SQL tokens"], split_rows)
    para(
        f"Combined gold examples analysed: {sql_eda['n_examples_total']}. "
        f"Unique databases referenced: {all_split.get('n_unique_dbs', '—')}."
    )

    # 7
    heading("7. SQL Query Complexity Analysis")
    fig("sql_complexity.png", "Figure 2: SQL feature prevalence")
    feat_defs = [
        ("has_where", "WHERE clause"),
        ("has_aggregation", "Aggregation"),
        ("has_join", "JOIN"),
        ("has_group_by", "GROUP BY"),
        ("has_order_by", "ORDER BY"),
        ("has_subquery", "Subquery"),
        ("has_limit", "LIMIT"),
        ("has_distinct", "DISTINCT"),
        ("has_having", "HAVING"),
        ("has_like", "LIKE"),
        ("has_between", "BETWEEN"),
        ("has_in", "IN"),
        ("has_intersect", "INTERSECT"),
        ("has_except", "EXCEPT"),
        ("has_union", "UNION"),
    ]
    split_names = [n for n, st in sql_eda["splits"].items() if n != "all" and st.get("n_examples")]
    if not split_names:
        split_names = ["all"]
    headers = ["SQL Feature"] + [f"{n.title()} %" for n in split_names]
    feature_rows = []
    for feat, label in feat_defs:
        row = [label]
        for name in split_names:
            row.append(sql_eda["splits"][name]["sql_complexity_pct"].get(feat, 0))
        feature_rows.append(row)
    table(headers, feature_rows)

    heading("7.1 Aggregation Function Breakdown", 2)
    table(["Function", "Count"], list(sql_eda["aggregations"].items()))

    heading("7.2 Query Complexity Tiers", 2)
    fig("complexity_tiers.png", "Figure 3: Complexity tiers")
    tier_rows = []
    for tk, label in [("simple", "Simple"), ("moderate", "Moderate"),
                      ("complex", "Complex"), ("extra", "Extra Hard")]:
        row = [label]
        for name, st in sql_eda["splits"].items():
            if name == "all" or not st.get("n_examples"):
                continue
            t = st["complexity_tiers"][tk]
            row.append(f"{t['count']} ({t['pct']}%)")
        tier_rows.append(row)
    th = ["Tier"] + [n.title() for n, st in sql_eda["splits"].items() if n != "all" and st.get("n_examples")]
    table(th or ["Tier", "All"], tier_rows)

    heading("7.3 JOIN Distribution", 2)
    fig("join_distribution.png", "Figure 4: JOINs per query")
    jd = sql_eda["join_count_distribution"]
    total_j = sum(jd.values()) or 1
    parts = ", ".join(f"{k} JOINs: {v} ({100 * v / total_j:.0f}%)" for k, v in jd.items())
    para(f"Distribution: {parts}.")

    # 8 — no NL questions in gold SQL
    heading("8. Natural Language Question Analysis")
    para(
        "Local gold SQL files do not include natural-language questions. "
        "Question-pattern EDA requires train/dev JSON; this run substitutes SQL hardness as a structural proxy."
    )
    fig("question_patterns.png", "Figure 5: Structural hardness proxy (no NL available locally)")

    # 9 Schema
    heading("9. Database Schema Analysis")
    para(
        f"Across {schema['n_dbs']} schemas: avg {schema['tables']['mean']} tables/DB "
        f"(median {schema['tables']['median']}, max {int(schema['tables']['max'])}), "
        f"avg {schema['columns']['mean']} columns/DB "
        f"(median {schema['columns']['median']}, max {int(schema['columns']['max'])}), "
        f"avg {schema['foreign_keys']['mean']} foreign keys/DB."
    )
    fig("column_types.png", "Figure 6: Column data type distribution")
    total_ct = sum(schema["column_types"].values()) or 1
    table(["Type", "Count", "%"], [
        [k, v, f"{100 * v / total_ct:.1f}%"] for k, v in schema["column_types"].items()
    ])

    heading("9.1 Domain Distribution (Heuristic)", 2)
    table(["Domain", "DB Count"], list(eda["domain_distribution"].items()))

    # 10 Profiling
    heading("10. SQLite Database Profiling")
    profiles = eda["database_profiles"]
    rows_list = [p["rows"] for p in profiles]
    para(
        f"Profiled {dq['dbs_profiled']} databases "
        f"({dq['sqlite_found']} with readable SQLite files).\n"
        f"Tables/DB: min {min(p['tables'] for p in profiles)}, max {max(p['tables'] for p in profiles)}, "
        f"mean {statistics.mean(p['tables'] for p in profiles):.1f}.\n"
        f"Rows/DB: min {min(rows_list)}, max {max(rows_list):,}, mean {statistics.mean(rows_list):,.0f}, "
        f"median {statistics.median(rows_list):.0f}."
    )
    fig("db_size_distribution.png", "Figure 7: Database row counts")
    fig("tables_per_db.png", "Figure 7b: Tables per database")
    fig("rows_per_table.png", "Figure 7c: Rows per table")
    fig("columns_per_table.png", "Figure 7d: Columns per table")
    top_rows = sorted(profiles, key=lambda x: -x["rows"])[:8]
    table(["Database", "Tables", "Rows", "Columns", "NULL %", "FKs"], [
        [p["db_id"], p["tables"], f"{p['rows']:,}", p["columns"], p["null_pct"], p["fk_count"]]
        for p in top_rows
    ])

    # 11 Quality
    heading("11. Data Quality Assessment")
    bullets([
        f"SQLite files found/opened: {dq['sqlite_found']} / {dq['dbs_profiled']}",
        f"Mean NULL % per DB: {dq['mean_null_pct_per_db']}% (max {dq['max_null_pct_per_db']}%)",
        f"Columns with NULLs: {dq['columns_with_nulls']} / {dq['total_columns_profiled']}",
        f"Empty databases (0 rows): {dq['dbs_with_zero_rows']}",
        f"Tables profiled: {dq['tables_profiled']}",
        "No remote downloads were used; analysis is fully offline from this directory.",
    ])

    # 12 samples
    heading("12. Sample Gold SQL Examples")
    samples = sql_eda["per_query_rows"][:6]
    table(["Database", "Tier", "Hardness", "JOINs", "SQL (truncated)"], [
        [r["db_id"], r["tier"], r["hardness"], r["n_joins"],
         r["query"][:80] + ("..." if len(r["query"]) > 80 else "")]
        for r in samples
    ])

    # 13
    heading("13. Text-to-SQL Pipeline Design Recommendations")
    for k, v in [
        ("Schema input", "Serialize tables/columns/types/FKs from schema.json (+ SQLite PRAGMA for runtime)."),
        ("JOIN modelling", f"{tc.get('has_join', 0)}% of local gold queries use JOIN; focus on INNER JOIN + FK paths."),
        ("FK supervision", f"{sql_eda['fk_alignment']['pct']}% of parsed JOIN ON pairs align with schema FKs."),
        ("Evaluation", "Prefer execution accuracy on SQLite over string match."),
        ("Exports", "Use eda_csvs/table_analysis.csv and join_analysis CSVs for debugging schema linking."),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)

    # 14
    heading("14. Preprocessing Completed")
    bullets([
        f"Profiled {dq['sqlite_found']} SQLite databases",
        "Exported database_summary.csv, table_analysis.csv, column_summary.csv",
        "Exported join_analysis.csv, join_conditions.csv, table_cooccurrence.csv, schema_fk_graph.csv, query_relations.csv",
        f"Classified {sql_eda['n_examples_total']} gold SQL queries",
        f"Generated {len(chart_files)} charts under figures/",
    ])

    # 15
    heading("15. Data Governance")
    bullets([
        "Source: Local Spider snapshot in this directory",
        "License: Academic/research use (original Spider terms)",
        "Privacy: Synthetic / anonymized research DBs",
        "Reproducibility: python generate_milestone2_report.py (or eda_spider.py)",
    ])

    # 16
    heading("16. Deliverables")
    table(["Deliverable", "Description"], [
        ["Deep SQL EDA Report.docx", "This EDA report (regenerated)"],
        ["eda_results.json", "Machine-readable statistics"],
        ["eda_csvs/database_summary.csv", "Per-database profiles"],
        ["eda_csvs/table_analysis.csv", "Per-table profiles"],
        ["eda_csvs/column_summary.csv", "Per-column NULL/cardinality"],
        ["eda_csvs/query_relations.csv", "Per-query SQL structure"],
        ["eda_csvs/join_analysis.csv", "Per-DB JOIN rates"],
        ["eda_csvs/join_conditions.csv", "Parsed JOIN ON pairs + FK flags"],
        ["eda_csvs/table_cooccurrence.csv", "Table pair co-occurrence"],
        ["eda_csvs/schema_fk_graph.csv", "Schema FK utilization"],
        ["figures/", f"{len(chart_files)} visualization charts"],
        ["generate_milestone2_report.py", "Reproducible EDA generator"],
    ])

    # 17
    heading("17. Planned Activities for Milestone 3")
    bullets([
        "Add full train/dev/test JSON if available for NL question-pattern EDA",
        "Implement schema serialization for model prompts",
        "Build execution-based evaluator on local SQLite DBs",
        "Train/evaluate baseline Text-to-SQL model",
        "Error analysis by JOIN / aggregation / wrong-table categories",
    ])

    # 18
    heading("18. Summary")
    para(
        f"Offline EDA covered {schema['n_dbs']} schemas and {sql_eda['n_examples_total']} gold SQL queries. "
        f"JOIN prevalence is {tc.get('has_join', 0)}%, aggregations {tc.get('has_aggregation', 0)}%, "
        f"subqueries {tc.get('has_subquery', 0)}%. "
        f"FK alignment on parsed JOIN conditions is {sql_eda['fk_alignment']['pct']}%. "
        f"Mean NULL rate across DBs is {dq['mean_null_pct_per_db']}%. "
        f"CSV exports under eda_csvs/ capture table and JOIN findings for Milestone 3."
    )

    # Appendix — Deep SQL EDA
    heading("Appendix: Deep SQL EDA Findings")
    para(
        "Advanced structural analysis of local gold SQL: JOIN types, hardness, subqueries, "
        "WHERE operators, table references, FK alignment, and relation topology."
    )

    heading("1. Overview", 2)
    para(
        f"Analysed {sql_eda['n_examples_total']} SQL queries. "
        f"JOIN count distribution: {sql_eda['join_count_distribution']}."
    )

    heading("2. JOIN Analysis", 2)
    fig("deep_join_count.png", "Figure A1: JOIN count distribution")
    join_pct = tc.get("has_join", 0)
    multi = sum(v for k, v in sql_eda["join_count_distribution"].items() if k >= 2)
    para(
        f"{join_pct}% of queries contain at least one JOIN. "
        f"{multi} queries have 2+ JOINs."
    )

    heading("2.1 JOIN Types", 3)
    fig("deep_join_types.png", "Figure A2: JOIN types")
    table(["JOIN type", "Count"], list(sql_eda["join_types"].items()))
    para("Spider typically uses bare JOIN (INNER). LEFT/RIGHT/CROSS are rare or absent.")

    heading("2.2 Databases with Highest JOIN Rates", 3)
    fig("deep_question_join_rate.png", "Figure A3: JOIN rate by database")
    table(["Database", "Queries", "With JOIN", "JOIN %"], [
        [r["db_id"], r["n_queries"], r["n_with_join"], r["join_rate_pct"]]
        for r in sql_eda["db_join_rates"][:12]
    ])

    heading("3. Subquery Analysis", 2)
    fig("deep_subquery_types.png", "Figure A4: Subquery locations")
    para(f"Subquery prevalence: {tc.get('has_subquery', 0)}%. Locations: {sql_eda['subquery_locations']}.")

    heading("4. Spider Official SQL Hardness", 2)
    fig("deep_spider_hardness.png", "Figure A5: Hardness distribution")
    fig("deep_hardness_splits.png", "Figure A6: Hardness by split")
    table(["Hardness", "Count"], list(sql_eda["hardness"].items()))

    heading("5. Table References per Query", 2)
    fig("deep_table_refs.png", "Figure A7: Tables referenced")
    table(["# Tables", "Queries"], list(sql_eda["table_ref_distribution"].items()))

    heading("6. WHERE Clause Analysis", 2)
    fig("deep_where_operators.png", "Figure A8: WHERE operators")
    table(["Operator", "Queries containing it"], list(sql_eda["where_operators"].items()))

    heading("7. Aggregation Analysis", 2)
    fig("deep_aggregations.png", "Figure A9: Aggregations")
    table(["Function", "Count"], list(sql_eda["aggregations"].items()))

    heading("8. SELECT / Other Clauses", 2)
    fig("deep_clause_prevalence.png", "Figure A10: Clause prevalence")

    heading("9. Query Relations Analysis", 2)
    heading("9.1 Relation Topology", 3)
    fig("deep_relation_topology.png", "Figure A11: Topology")
    table(["Topology", "Count"], list(sql_eda["topology"].items()))

    heading("9.2 Foreign Key Alignment in JOINs", 3)
    fig("deep_fk_alignment.png", "Figure A12: FK alignment")
    fa = sql_eda["fk_alignment"]
    para(
        f"Of {fa['join_conditions']} parsed JOIN ON column pairs, "
        f"{fa['fk_aligned']} ({fa['pct']}%) match schema foreign keys."
    )

    heading("9.3 Table Co-occurrence & JOIN Patterns", 3)
    fig("deep_table_pairs.png", "Figure A13: Top table pairs")
    table(["Database", "Table A", "Table B", "Co-occurrence"], [
        [r["db_id"], r["table_a"], r["table_b"], r["cooccurrence"]]
        for r in sql_eda["table_pair_rows"][:15]
    ])

    heading("9.4 Schema FK Utilization by Database", 3)
    fk_util = sorted(eda["schema_fk_graph"], key=lambda r: -r["n_foreign_keys"])[:12]
    table(["Database", "Tables", "FKs", "FK coverage %"], [
        [r["db_id"], r["n_tables"], r["n_foreign_keys"], r["fk_coverage_pct"]] for r in fk_util
    ])

    heading("10. Key Takeaways for Text-to-SQL", 2)
    bullets([
        f"{join_pct}% of local gold queries need JOINs — schema linking is mandatory.",
        f"JOIN types observed: {sql_eda['join_types'] or 'none'} — model INNER JOIN first.",
        f"{fa['pct']}% of JOIN ON pairs align with schema FKs — use FKs as join supervision.",
        f"Subqueries in {tc.get('has_subquery', 0)}% of queries — mostly WHERE/SELECT nested SELECTs.",
        f"Dominant aggregation: {(list(sql_eda['aggregations'].keys()) or ['n/a'])[0]}.",
        "Evaluate with execution accuracy so equivalent JOIN order still passes.",
        "See eda_csvs/ for full per-table and per-JOIN exports.",
    ])

    out = ROOT / "Deep SQL EDA Report.docx"
    doc.save(out)
    print(f"Saved {out}")
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_eda() -> tuple[dict, dict]:
    schemas = load_schemas()
    examples = load_gold_sql_files()
    if not examples:
        print("WARNING: No *_gold.sql files found — SQL/JOIN sections will be empty.")

    print(f"Loaded {len(schemas)} schemas, {len(examples)} gold SQL examples")
    schema_index = build_schema_index(schemas)
    print("Profiling SQLite databases (this may take a few minutes)...")
    db_profiles, table_rows, column_rows = profile_databases(schemas)
    schema_fk_graph = analyze_schema_joins(schema_index)
    sql_eda = analyze_sql_examples(examples, schema_index)

    null_pcts = [p["null_pct"] for p in db_profiles]
    eda = {
        "overview": {
            "dataset": "Spider (local snapshot)",
            "task": "Text-to-SQL",
            "dialect": "SQLite",
            "n_schemas": len(schemas),
            "n_gold_sql": len(examples),
        },
        "inputs": {
            "n_schemas": len(schemas),
            "n_sqlite": sum(1 for p in db_profiles if p["sqlite_found"]),
            "gold_files": [p.name for p in sorted(ROOT.glob("*_gold.sql"))],
        },
        "schema_stats": schema_level_stats(schemas),
        "database_profiles": db_profiles,
        "table_analysis": table_rows,
        "domain_distribution": dict(Counter(p["domain"] for p in db_profiles).most_common()),
        "schema_fk_graph": schema_fk_graph,
        "data_quality": {
            "dbs_profiled": len(db_profiles),
            "sqlite_found": sum(1 for p in db_profiles if p["sqlite_found"]),
            "dbs_with_zero_rows": sum(1 for p in db_profiles if p["rows"] == 0),
            "mean_null_pct_per_db": round(statistics.mean(null_pcts), 2) if null_pcts else 0,
            "max_null_pct_per_db": max(null_pcts) if null_pcts else 0,
            "columns_with_nulls": sum(1 for r in column_rows if r["null_count"] > 0),
            "total_columns_profiled": len(column_rows),
            "tables_profiled": len(table_rows),
        },
    }

    # ---- CSV exports ----
    write_csv(CSV_DIR / "database_summary.csv", db_profiles)
    write_csv(CSV_DIR / "table_analysis.csv", table_rows)
    write_csv(CSV_DIR / "column_summary.csv", column_rows)
    write_csv(CSV_DIR / "schema_fk_graph.csv", schema_fk_graph)
    write_csv(CSV_DIR / "query_relations.csv", sql_eda["per_query_rows"])
    write_csv(CSV_DIR / "join_analysis.csv", sql_eda["db_join_rates"])
    write_csv(CSV_DIR / "join_conditions.csv", sql_eda["join_condition_rows"])
    write_csv(CSV_DIR / "table_cooccurrence.csv", sql_eda["table_pair_rows"])

    # Also mirror key CSVs at ROOT for backwards compatibility
    write_csv(ROOT / "database_summary.csv", db_profiles)
    write_csv(ROOT / "column_summary.csv", column_rows)

    # Compact JSON (drop huge row lists)
    eda_json = {
        **eda,
        "database_profiles": db_profiles,
        "table_analysis_preview": table_rows[:20],
        "sql_summary": {
            k: v for k, v in sql_eda.items()
            if k not in ("per_query_rows", "join_condition_rows", "table_pair_rows")
        },
        "sql_splits": sql_eda["splits"],
    }
    # Remove heavy nested copies from table_analysis full list in json — keep summary only
    eda_json.pop("table_analysis", None)

    with open(ROOT / "eda_results.json", "w") as f:
        json.dump(eda_json, f, indent=2, default=str)

    return eda, sql_eda


def main():
    print("Running Spider EDA from local files...")
    eda, sql_eda = build_eda()
    print("Generating charts...")
    charts = make_charts(eda, sql_eda)
    print("Writing Deep SQL EDA Report.docx...")
    write_docx(eda, sql_eda, charts)
    print("\nDone. Deliverables:")
    for f in [
        "Deep SQL EDA Report.docx", "eda_results.json",
        "database_summary.csv", "column_summary.csv",
    ]:
        print(f"  - {ROOT / f}")
    print(f"  - {CSV_DIR}/")
    for p in sorted(CSV_DIR.glob("*.csv")):
        print(f"      {p.name}")
    print(f"  - {FIGURES}/ ({len(charts)} charts)")


if __name__ == "__main__":
    main()
